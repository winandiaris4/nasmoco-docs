"""
generator.py — Core document generation logic
Markdown → HTML → PDF / DOCX / HTML
"""
import io
import re
import requests
from datetime import date
from pathlib import Path

import mistune
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML as WeasyprintHTML
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .config import get_config

TEMPLATE_DIR = Path(__file__).parent / "templates"
STYLES_DIR = TEMPLATE_DIR / "styles"
LAYOUTS_DIR = TEMPLATE_DIR / "layouts"

jinja_env = Environment(
    loader=FileSystemLoader(str(TEMPLATE_DIR)),
    autoescape=False,
)


# ─────────────────────────────────────────────
# Markdown Parser
# ─────────────────────────────────────────────

def parse_markdown(source_path: Path) -> dict:
    """
    Parse Markdown file menjadi structured content.
    Mengekstrak judul, metadata (dari baris | Key | Value |), dan body HTML.
    Menghilangkan duplikasi judul dan tabel metadata dari body.
    """
    text = source_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    metadata = {}
    content_lines = []
    
    # Kumpulkan metadata dan hapus baris-baris tersebut agar tidak dobel di body
    in_metadata_table = False
    first_h1_stripped = False

    for line in lines:
        stripped = line.strip()
        
        # Hapus H1 pertama (karena sudah dijadikan judul dokumen di template)
        if stripped.startswith("# ") and not first_h1_stripped:
            first_h1_stripped = True
            continue
            
        # Deteksi awal tabel metadata
        if stripped.startswith("|") and ("**Dari" in stripped or "**Kepada" in stripped or "**Nomor" in stripped or "**Tanggal" in stripped or "**Revisi" in stripped):
            in_metadata_table = True
            # Ekstrak data
            parts = [p.strip() for p in line.split("|") if p.strip()]
            if len(parts) == 2:
                key = parts[0].replace("**", "").strip()
                val = parts[1].replace("**", "").strip()
                metadata[key] = val
            continue
            
        # Deteksi baris pemisah tabel metadata
        if in_metadata_table and (stripped.startswith("|---|") or stripped.startswith("| :---") or stripped.startswith("|---")):
            continue
            
        # Deteksi baris-baris isi tabel metadata lainnya
        if in_metadata_table and stripped.startswith("|"):
            parts = [p.strip() for p in line.split("|") if p.strip()]
            if len(parts) == 2:
                key = parts[0].replace("**", "").strip()
                val = parts[1].replace("**", "").strip()
                metadata[key] = val
            continue
            
        # Jika tabel metadata selesai (bertemu baris kosong atau garis pembatas)
        if in_metadata_table and not stripped.startswith("|"):
            in_metadata_table = False
            # Hapus hr pemisah teratas jika ada
            if stripped == "---":
                continue
            
        content_lines.append(line)

    body_md = "\n".join(content_lines)
    body_html = mistune.html(body_md)

    # Temukan judul dari baris asli pertama heading #
    title = ""
    for line in lines:
        if line.startswith("# "):
            title = line[2:].strip()
            break

    return {
        "title": title,
        "metadata": metadata,
        "body_html": body_html,
        "source_path": str(source_path),
        "generated_date": date.today().strftime("%d %B %Y"),
    }


# ─────────────────────────────────────────────
# Logo Fetcher
# ─────────────────────────────────────────────

def fetch_logo_base64(logo_url: str) -> str:
    """Fetch logo dari URL atau path lokal dan encode ke base64 untuk embed di HTML."""
    import base64
    import mimetypes
    try:
        # Jika berupa url
        if logo_url.startswith("http://") or logo_url.startswith("https://"):
            resp = requests.get(logo_url, timeout=10)
            resp.raise_for_status()
            content = resp.content
            content_type = resp.headers.get("Content-Type", "image/png")
        else:
            # Cari path lokal (relatif terhadap nasmoco-docs root)
            p = Path(logo_url)
            if not p.is_absolute():
                from .config import REPO_ROOT
                p = REPO_ROOT / logo_url
            
            if p.exists():
                content = p.read_bytes()
                content_type, _ = mimetypes.guess_type(str(p))
                content_type = content_type or "image/png"
            else:
                return ""

        b64 = base64.b64encode(content).decode("utf-8")
        return f"data:{content_type};base64,{b64}"
    except Exception:
        return ""  # Fallback: tampilkan teks nama brand saja


# ─────────────────────────────────────────────
# HTML Generator
# ─────────────────────────────────────────────

def generate_html(
    source_path: Path,
    doc_type: str,
    style: str = "modern",
    layout: str = "table",
    overrides: dict = None,
) -> str:
    """Render Markdown ke HTML dengan kop surat Winamus."""
    cfg = get_config()
    content = parse_markdown(source_path)
    logo_data = fetch_logo_base64(cfg["company"]["logo_url"])

    # Load CSS style
    style_file = STYLES_DIR / f"{style}.css"
    css_content = style_file.read_text(encoding="utf-8") if style_file.exists() else ""

    # Pilih layout template
    template_name = f"layouts/{doc_type}_{layout}.html"
    # Fallback ke generic jika layout spesifik tidak ada
    if not (LAYOUTS_DIR / f"{doc_type}_{layout}.html").exists():
        template_name = "layouts/generic.html"

    template = jinja_env.get_template(template_name)

    context = {
        "company": cfg["company"],
        "colors": cfg["colors"],
        "css": css_content,
        "logo_data": logo_data,
        "doc_type": doc_type,
        "style": style,
        "layout": layout,
        **content,
        **(overrides or {}),
    }

    return template.render(**context)


# ─────────────────────────────────────────────
# PDF Generator
# ─────────────────────────────────────────────

def generate_pdf(html_content: str) -> bytes:
    """Convert HTML ke PDF menggunakan WeasyPrint."""
    pdf_bytes = WeasyprintHTML(string=html_content).write_pdf()
    return pdf_bytes


# ─────────────────────────────────────────────
# DOCX Generator
# ─────────────────────────────────────────────

def generate_docx(
    source_path: Path,
    doc_type: str,
    style: str = "modern",
    overrides: dict = None,
) -> bytes:
    """Generate DOCX dengan python-docx. Lebih sederhana dari PDF — plain structure."""
    cfg = get_config()
    content = parse_markdown(source_path)

    doc = Document()

    # ── Page margins ──
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    # ── Header: Nama perusahaan ──
    header = doc.add_heading(cfg["company"]["name"], level=1)
    header.runs[0].font.size = Pt(16)
    header.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Tagline & info
    info = doc.add_paragraph(
        f"{cfg['company']['website']} | {cfg['company']['email']} | {cfg['company']['phone']}"
    )
    info.runs[0].font.size = Pt(9)
    info.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph("─" * 80)

    # ── Judul dokumen ──
    doc.add_heading(content["title"], level=2)

    # ── Body: parse HTML sederhana ke paragraf ──
    # Strip HTML tags untuk DOCX plain text
    clean_text = re.sub(r"<[^>]+>", "", content["body_html"])
    clean_text = re.sub(r"\n{3,}", "\n\n", clean_text).strip()

    for para in clean_text.split("\n\n"):
        para = para.strip()
        if para:
            p = doc.add_paragraph(para)
            p.runs[0].font.size = Pt(10.5) if p.runs else None

    # ── Footer ──
    doc.add_paragraph("─" * 80)
    footer_p = doc.add_paragraph(
        f"© {date.today().year} {cfg['company']['name']} — {cfg['company']['address']}"
    )
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ─────────────────────────────────────────────
# Main Entry Point
# ─────────────────────────────────────────────

def generate(
    source_path: str | Path,
    doc_type: str,
    style: str = "modern",
    layout: str = "table",
    formats: list[str] = None,
    output_dir: str | Path = None,
    project_docs_dir: str | Path = None,
    overrides: dict = None,
) -> dict[str, Path]:
    """
    Generate dokumen dari Markdown ke format yang diminta.

    Args:
        source_path: Path ke file Markdown sumber
        doc_type: Tipe dokumen (quotation, invoice, dll)
        style: Visual style (modern, classic, bold)
        layout: Layout varian (table, narrative, dll)
        formats: List format output ["pdf", "docx", "html"]
        output_dir: Folder output utama (default: generator/output/)
        project_docs_dir: Jika diisi, hasil juga disimpan ke sini (project/xxx/Docs/)
        overrides: Data tambahan untuk di-inject ke template

    Returns:
        Dict berisi {format: Path} untuk setiap file yang dihasilkan
    """
    cfg = get_config()
    source_path = Path(source_path)
    formats = formats or ["pdf", "docx", "html"]

    # Tentukan output dir
    if output_dir is None:
        output_dir = Path(__file__).parent.parent / cfg["defaults"]["output_dir"]
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Nama file output berdasarkan nama source
    stem = source_path.stem
    results = {}

    # Generate HTML dulu (dipakai semua format)
    html_content = generate_html(source_path, doc_type, style, layout, overrides)

    for fmt in formats:
        filename = f"{stem}_{style}.{fmt}"
        out_path = output_dir / filename

        if fmt == "html":
            out_path.write_text(html_content, encoding="utf-8")
        elif fmt == "pdf":
            pdf_bytes = generate_pdf(html_content)
            out_path.write_bytes(pdf_bytes)
        elif fmt == "docx":
            docx_bytes = generate_docx(source_path, doc_type, style, overrides)
            out_path.write_bytes(docx_bytes)

        results[fmt] = out_path

        # Jika project_docs_dir diisi, salin juga ke sana
        if project_docs_dir:
            project_docs_dir = Path(project_docs_dir)
            project_docs_dir.mkdir(parents=True, exist_ok=True)
            dest = project_docs_dir / filename
            dest.write_bytes(out_path.read_bytes())

    return results
